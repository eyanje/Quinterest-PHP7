from docx import Document
from PyPDF2 import PdfFileReader

import io, os
import urllib.request, urllib.parse

# This module reads text from docx and pdf files

def read_docx(path):
    # path = urllib.parse.quote(path)
    try:
        if (path.startswith('C:')):
            with open(path, 'rb') as file:
                document = Document(file)
        else:
            with urllib.request.urlopen(path) as file:
                document = Document(io.BytesIO(file.read()))
    except urllib.error.HTTPError as error:
        print(f'Error loading docx {path}: {error}')
        return None
    except ValueError:
        with open(path, 'rb') as file:
            document = Document(file)
    except FileNotFoundError:
        print(f'{path} not found!')
        document = None
    if document != None:
        return document.paragraphs

def read_pdf(path):
    # path = urllib.parse.quote(path)
    try:
        with urllib.request.urlopen(urllib.request.Request(path, headers={'User-Agent':' Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:67.0) Gecko/20100101 Firefox/67.0'})) as file:
            pdf = PdfFileReader(io.BytesIO(file.read()))
            for page_num in range(pdf.numPages):
                page = pdf.getPage(page_num)
                print(page.extractText())
    except urllib.error.HTTPError as error:
        print(f'Error loading pdf {path}: {error}')

def extract_texts(paragraphs):
    texts = []
    if (paragraphs == None):
        return None
    else:
        for paragraph in paragraphs:
            if (len(paragraph.text) > 2):
                texts.append(paragraph.text)
    return texts

def frequency_dict(paragraphs):
    frequencies = {}
    for paragraph in paragraphs:
        for word in paragraph.text.split():
            if (word not in frequencies):
                frequencies[word] = 0
            frequencies[word] += 1
    return frequencies
